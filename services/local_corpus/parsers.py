from __future__ import annotations

from pathlib import Path

from core.knowledge_models import SourceKind
from core.local_corpus_models import ParsedLocalDocument, TableProfile

from .table_profiler import profile_csv, profile_rows


class ParserRegistry:
    parser_version = "local-corpus-v1"

    def parse(self, source_id: str, path: str | Path) -> ParsedLocalDocument:
        file_path = Path(path)
        suffix = file_path.suffix.lower()
        if suffix == ".md":
            return ParsedLocalDocument(
                source_id=source_id,
                source_kind=SourceKind.MARKDOWN,
                markdown_text=self._read_text(file_path),
            )
        if suffix == ".txt":
            return ParsedLocalDocument(
                source_id=source_id,
                source_kind=SourceKind.TXT,
                markdown_text=self._read_text(file_path),
            )
        if suffix == ".csv":
            profile = profile_csv(source_id, file_path)
            return ParsedLocalDocument(
                source_id=source_id,
                source_kind=SourceKind.CSV,
                markdown_text=profile.summary_markdown,
                table_profiles=[profile],
            )
        if suffix == ".docx":
            return self._parse_docx(source_id, file_path)
        if suffix == ".pdf":
            return self._parse_pdf(source_id, file_path)
        if suffix == ".xlsx":
            return self._parse_xlsx(source_id, file_path)
        return ParsedLocalDocument(
            source_id=source_id,
            source_kind=SourceKind.UNKNOWN,
            markdown_text="",
            metadata={"parser_error": f"unsupported extension: {suffix}"},
        )

    def _read_text(self, path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "cp949"):
            try:
                return path.read_text(encoding=encoding).strip()
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="utf-8", errors="replace").strip()

    def _parse_docx(self, source_id: str, path: Path) -> ParsedLocalDocument:
        try:
            import docx  # type: ignore
        except Exception as exc:
            raise RuntimeError("python-docx is required to parse .docx files") from exc
        document = docx.Document(str(path))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        return ParsedLocalDocument(
            source_id=source_id,
            source_kind=SourceKind.DOCX,
            markdown_text="\n\n".join(paragraphs),
        )

    def _parse_pdf(self, source_id: str, path: Path) -> ParsedLocalDocument:
        reader_cls = None
        try:
            from pypdf import PdfReader  # type: ignore

            reader_cls = PdfReader
        except Exception:
            try:
                from PyPDF2 import PdfReader  # type: ignore

                reader_cls = PdfReader
            except Exception as exc:
                raise RuntimeError("pypdf or PyPDF2 is required to parse .pdf files") from exc
        reader = reader_cls(str(path))
        pages: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"## Page {index}\n\n{text}")
        return ParsedLocalDocument(
            source_id=source_id,
            source_kind=SourceKind.PDF,
            markdown_text="\n\n".join(pages),
        )

    def _parse_xlsx(self, source_id: str, path: Path) -> ParsedLocalDocument:
        try:
            import openpyxl  # type: ignore
        except Exception as exc:
            raise RuntimeError("openpyxl is required to parse .xlsx files") from exc

        workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        profiles: list[TableProfile] = []
        parts: list[str] = []
        for sheet in workbook.worksheets:
            rows: list[list[str]] = []
            for idx, row in enumerate(sheet.iter_rows(values_only=True)):
                rows.append([str(cell or "") for cell in row])
                if idx >= 200:
                    break
            profile = profile_rows(source_id, sheet.title, rows)
            profiles.append(profile)
            parts.append(profile.summary_markdown)
        try:
            workbook.close()
        except Exception:
            pass
        return ParsedLocalDocument(
            source_id=source_id,
            source_kind=SourceKind.XLSX,
            markdown_text="\n\n---\n\n".join(parts),
            table_profiles=profiles,
        )


__all__ = ["ParserRegistry"]
