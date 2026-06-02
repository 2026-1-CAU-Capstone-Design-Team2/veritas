# -*- coding: utf-8 -*-
"""Scenario A synthetic internal data generator (시나리오 A — 삼성전자 DS/HBM).

Generates the "민감 사내 데이터" files for the local-corpus integration demo,
designed against the ACTUAL web research output in runs/삼성전자/:

* Numbers that MATCH official web figures   -> verification "supports"
* Numbers that deliberately MISMATCH        -> crosscheck "numeric_mismatch" flags
* Internal-only data that fills the web research's Core Gaps:
    - Core Gap 1: HBM 경쟁사별 점유율 / 생산능력 / 포트폴리오   (PDF + XLSX)
    - Core Gap 2: 2026년 HBM 수급/가격(ASP) 전망               (XLSX)
    - Core Gap 3: DS부문 원가/재고충당/공정별 비용 구조          (DOCX + CSV)

Outputs (test_data/scenario_a/):
    DS부문_2025년_4분기_내부결산보고_대외비.docx
    HBM4_경쟁구도_내부평가_대외비.pdf            (Word COM 변환; 실패 시 docx 유지)
    HBM_고객사별_수주현황_2026_대외비.xlsx
    DS부문_월별_사업부실적_2025.csv

Plus a tester-only answer sheet OUTSIDE the local-access folder:
    test_data/scenario_a_expected/expected_values.json

Usage:
    python tests/fixtures/generate_scenario_data.py
"""
from __future__ import annotations

import csv
import json
import sys
from pathlib import Path

# 한국어 Windows 콘솔(cp949)에서 em-dash 등 특수문자 print가 죽지 않도록 보호.
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(errors="replace")

REPO_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = REPO_ROOT / "test_data" / "scenario_a"
EXPECTED_DIR = REPO_ROOT / "test_data" / "scenario_a_expected"

# ---------------------------------------------------------------------------
# 수치 설계 (웹 조사 결과 runs/삼성전자/ 기준)
#
#   일치 (supports)                  불일치 (crosscheck 탐지 대상)
#   ------------------------------   ------------------------------------------
#   4Q25 연결 매출 93.8조원           DS 4Q 영업이익: 내부 15.8조 (웹 16.4조)
#   4Q25 연결 영업이익 20.1조원       연간 R&D: 내부 36.2조 (웹 37.7조)
#   DS 4Q 매출 44조원                 SK하이닉스 HBM4 점유율: 내부 62% (웹 70%)
#   DS 이익 비중 82%                  자사 HBM4 수율: 내부 68% (웹 보도 50~60%)
#   HBM4 시장 546억 달러              마이크론 점유율: 내부 15% (웹 22%)
#   HBM4 대역폭 2.8TB/s
#   DS 설비투자 47.5조원
#   SK하이닉스 M16 월 4만장
# ---------------------------------------------------------------------------

# CSV: 4Q 합계가 공식 발표(매출 44조)와 내부 관리회계(영업이익 15.8조)에 정확히 맞도록 설계
MONTHLY_REVENUE_TARGET = {  # 단위: 억원 (DS부문 전체)
    "2025-01": 98_000, "2025-02": 95_000, "2025-03": 104_000,   # 1Q: 29.7조
    "2025-04": 88_000, "2025-05": 84_000, "2025-06": 90_000,    # 2Q: 26.2조 (부진 — 웹 서사와 일치)
    "2025-07": 108_000, "2025-08": 113_000, "2025-09": 121_000,  # 3Q: 34.2조
    "2025-10": 140_000, "2025-11": 145_000, "2025-12": 155_000,  # 4Q: 44.0조 ← 웹 공식 발표 일치
}
MONTHLY_PROFIT_TARGET = {  # 단위: 억원
    "2025-01": 21_000, "2025-02": 19_500, "2025-03": 23_500,    # 1Q: 6.4조
    "2025-04": 9_500, "2025-05": 8_000, "2025-06": 10_500,      # 2Q: 2.8조 (재고충당 영향)
    "2025-07": 27_000, "2025-08": 30_000, "2025-09": 33_000,    # 3Q: 9.0조
    "2025-10": 49_000, "2025-11": 52_000, "2025-12": 57_000,    # 4Q: 15.8조 ← 내부 관리회계 (웹 16.4조와 불일치)
}

# 사업부 × 제품군 구조와 매출 배분 비율
BUSINESS_LINES = [
    # (사업부, 제품군, 매출 비중, 이익 기여 비중)
    ("메모리", "HBM", 0.28, 0.46),
    ("메모리", "DRAM", 0.30, 0.38),
    ("메모리", "NAND", 0.18, 0.20),
    ("파운드리", "선단공정(2nm/3nm)", 0.12, -0.028),
    ("파운드리", "성숙공정", 0.07, -0.012),
    ("시스템LSI", "SoC(엑시노스)", 0.025, -0.006),
    ("시스템LSI", "이미지센서", 0.020, 0.004),
    ("시스템LSI", "DDI", 0.005, 0.002),
]
REGIONS = [("북미", 0.40), ("중국", 0.22), ("국내", 0.15), ("유럽", 0.13), ("기타", 0.10)]

# XLSX 수주 데이터: 고객사 풀 (실존 기업명 — 웹 기사와 token overlap을 위해 의도적 사용)
CUSTOMERS = [
    ("NVIDIA", "북미", 0.38),
    ("AMD", "북미", 0.14),
    ("Broadcom", "북미", 0.09),
    ("Google", "북미", 0.10),
    ("Amazon AWS", "북미", 0.08),
    ("Microsoft Azure", "북미", 0.07),
    ("Meta", "북미", 0.05),
    ("Tesla", "북미", 0.03),
    ("Alibaba Cloud", "중국", 0.03),
    ("아틀라스컴퓨팅", "국내", 0.02),
    ("퀀텀리프AI", "국내", 0.01),
]
PRODUCTS = [
    # (제품명, 개당 단가 USD, 2025 비중, 2026 비중)
    ("HBM3E 8단", 95, 0.40, 0.15),
    ("HBM3E 12단", 142, 0.45, 0.30),
    ("HBM4 12단", 228, 0.12, 0.38),
    ("HBM4 16단", 312, 0.03, 0.17),
]
ORDER_MONTHS = [
    "2025-01", "2025-02", "2025-03", "2025-04", "2025-05", "2025-06",
    "2025-07", "2025-08", "2025-09", "2025-10", "2025-11", "2025-12",
    "2026-01", "2026-02", "2026-03", "2026-04", "2026-05",
]


def _check_deps() -> None:
    missing = []
    try:
        import docx  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    try:
        import openpyxl  # noqa: F401
    except ImportError:
        missing.append("openpyxl")
    if missing:
        sys.exit(f"누락된 패키지: {', '.join(missing)} — pip install 후 다시 실행하세요.")


# ---------------------------------------------------------------------------
# ① CSV — DS부문 월별 사업부 실적 (480행, Core Gap 3 + table_query 무손실 검증)
# ---------------------------------------------------------------------------

def generate_csv(path: Path) -> dict:
    rows: list[list] = []
    header = ["월", "사업부", "제품군", "지역", "매출(억원)", "영업이익(억원)", "원가율(%)", "재고충당금(억원)"]

    total_revenue = 0
    total_profit = 0
    q4_revenue = 0
    q4_profit = 0
    hbm_revenue = 0

    for month, month_revenue in MONTHLY_REVENUE_TARGET.items():
        month_profit = MONTHLY_PROFIT_TARGET[month]
        is_q2 = month in ("2025-04", "2025-05", "2025-06")
        is_q4 = month in ("2025-10", "2025-11", "2025-12")

        # 사업부×제품군별 월 목표를 비율로 배분한 뒤, 지역으로 다시 배분.
        # 마지막 항목에서 잔차를 보정해 월 합계가 목표와 정확히 일치하도록 한다.
        line_revenue_alloc: list[int] = []
        line_profit_alloc: list[int] = []
        for index, (_, _, rev_share, profit_share) in enumerate(BUSINESS_LINES):
            if index == len(BUSINESS_LINES) - 1:
                line_revenue_alloc.append(month_revenue - sum(line_revenue_alloc))
                line_profit_alloc.append(month_profit - sum(line_profit_alloc))
            else:
                line_revenue_alloc.append(int(month_revenue * rev_share))
                line_profit_alloc.append(int(month_profit * profit_share))

        for line_index, (division, product, _, _) in enumerate(BUSINESS_LINES):
            line_rev = line_revenue_alloc[line_index]
            line_profit = line_profit_alloc[line_index]
            # 원가율: 메모리는 낮고(고마진) 파운드리/LSI는 높음. 2Q는 재고충당으로 상승.
            if division == "메모리":
                cost_ratio = 56.0 + (3.5 if is_q2 else 0.0) + line_index * 1.1
            elif division == "파운드리":
                cost_ratio = 88.0 + (4.0 if is_q2 else 0.0)
            else:
                cost_ratio = 81.0 + (2.0 if is_q2 else 0.0)
            # 재고충당금: 2Q에 집중 (웹의 "2Q 재고자산 평가 충당금" 서사와 일치)
            provision_base = int(line_rev * (0.045 if is_q2 else 0.004))

            region_rev_alloc: list[int] = []
            region_profit_alloc: list[int] = []
            for region_index, (_, region_share) in enumerate(REGIONS):
                if region_index == len(REGIONS) - 1:
                    region_rev_alloc.append(line_rev - sum(region_rev_alloc))
                    region_profit_alloc.append(line_profit - sum(region_profit_alloc))
                else:
                    region_rev_alloc.append(int(line_rev * region_share))
                    region_profit_alloc.append(int(line_profit * region_share))

            for region_index, (region, region_share) in enumerate(REGIONS):
                region_rev = region_rev_alloc[region_index]
                region_profit = region_profit_alloc[region_index]
                region_provision = int(provision_base * region_share)
                rows.append([
                    month,
                    division,
                    product,
                    region,
                    f"{region_rev:,}",
                    f"{region_profit:,}",
                    f"{cost_ratio:.1f}",
                    f"{region_provision:,}",
                ])
                total_revenue += region_rev
                total_profit += region_profit
                if is_q4:
                    q4_revenue += region_rev
                    q4_profit += region_profit
                if product == "HBM":
                    hbm_revenue += region_rev

    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(header)
        writer.writerows(rows)

    return {
        "file": path.name,
        "row_count": len(rows),
        "total_revenue_억원": total_revenue,
        "total_profit_억원": total_profit,
        "q4_revenue_억원 (웹 공식 발표 44조원과 일치해야 함)": q4_revenue,
        "q4_profit_억원 (내부 관리회계 15.8조 — 웹 16.4조와 불일치 설계)": q4_profit,
        "hbm_연간매출_억원": hbm_revenue,
    }


# ---------------------------------------------------------------------------
# ② XLSX — HBM 고객사별 수주현황 + ASP 전망 + 경쟁사 생산능력 (Core Gap 1, 2)
# ---------------------------------------------------------------------------

def generate_xlsx(path: Path) -> dict:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    workbook = openpyxl.Workbook()
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")

    def style_header(sheet, column_count: int) -> None:
        for column in range(1, column_count + 1):
            cell = sheet.cell(row=1, column=column)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

    # -- 시트 1: 수주현황 (320행+) ------------------------------------------
    orders = workbook.active
    orders.title = "수주현황"
    orders.append([
        "수주번호", "수주일", "고객사", "지역", "제품", "수량(천개)",
        "단가(달러)", "수주액(백만달러)", "납기", "상태",
    ])

    order_rows = []
    order_number = 0
    total_order_musd = 0.0
    by_customer: dict[str, float] = {}
    by_product: dict[str, float] = {}
    hbm4_order_musd = 0.0

    for month_index, month in enumerate(ORDER_MONTHS):
        year = int(month[:4])
        is_2026 = year == 2026
        # 월별 수주 건수: 2025 초반 12건 → 2026년 28건 (HBM4 전환 가속)
        deals_this_month = 12 + month_index
        for deal_index in range(deals_this_month):
            order_number += 1
            customer, region, _ = CUSTOMERS[(order_number * 7 + deal_index) % len(CUSTOMERS)]
            # 제품 선택: 연도별 비중에 따라 결정론적으로 분포
            product_weights = [p[3] if is_2026 else p[2] for p in PRODUCTS]
            cumulative = 0.0
            selector = ((order_number * 13 + deal_index * 5) % 100) / 100.0
            product_name, unit_price = PRODUCTS[-1][0], PRODUCTS[-1][1]
            for (name, price, *_), weight in zip(PRODUCTS, product_weights):
                cumulative += weight
                if selector <= cumulative:
                    product_name, unit_price = name, price
                    break
            # 수량: 고객 규모에 따라 80천개 ~ 950천개
            quantity_k = 80 + ((order_number * 37 + deal_index * 11) % 870)
            amount_musd = round(quantity_k * 1000 * unit_price / 1_000_000, 1)
            day = (order_number * 3 + deal_index) % 27 + 1
            delivery_month_index = min(month_index + 3 + (deal_index % 4), len(ORDER_MONTHS) - 1)
            status = "완료" if month_index < 10 else ("진행중" if deal_index % 5 else "검수중")

            order_rows.append([
                f"PO-{year}-{order_number:04d}",
                f"{month}-{day:02d}",
                customer,
                region,
                product_name,
                quantity_k,
                unit_price,
                amount_musd,
                f"{ORDER_MONTHS[delivery_month_index]}-15",
                status,
            ])
            total_order_musd += amount_musd
            by_customer[customer] = round(by_customer.get(customer, 0.0) + amount_musd, 1)
            by_product[product_name] = round(by_product.get(product_name, 0.0) + amount_musd, 1)
            if product_name.startswith("HBM4"):
                hbm4_order_musd += amount_musd

    for row in order_rows:
        orders.append(row)
    style_header(orders, 10)
    for column, width in zip("ABCDEFGHIJ", (16, 12, 18, 8, 14, 12, 12, 18, 12, 10)):
        orders.column_dimensions[column].width = width

    # -- 시트 2: ASP 가격전망 (Core Gap 2 — 웹에 없는 내부 전용 데이터) -------
    asp = workbook.create_sheet("ASP_가격전망")
    asp.append(["분기", "HBM3E ASP(달러)", "HBM4 ASP(달러)", "HBM4 프리미엄(%)", "수급 전망", "비고"])
    asp_rows = [
        ["2025-Q1", 118, None, None, "HBM3E 공급 부족", "HBM4 양산 전"],
        ["2025-Q2", 121, None, None, "균형", "HBM4 고객 인증 진행"],
        ["2025-Q3", 124, 215, 73.4, "HBM4 공급 극히 제한", "초기 샘플 출하"],
        ["2025-Q4", 126, 222, 76.2, "HBM4 공급 부족", "양산 램프업"],
        ["2026-Q1", 122, 228, 86.9, "HBM4 공급 부족 지속", "엔비디아 루빈 수요 본격화"],
        ["2026-Q2", 118, 232, 96.6, "HBM4 타이트", "16단 비중 확대"],
        ["2026-Q3", 112, 226, 101.8, "균형 전환 예상", "경쟁사 공급 확대"],
        ["2026-Q4", 105, 218, 107.6, "공급 우위 전환 리스크", "가격 하락 압력 시작"],
    ]
    for row in asp_rows:
        asp.append(row)
    asp.append([])
    asp.append(["※ 내부 분석 기준 2026년 연평균 HBM4 가격 프리미엄은 HBM3E 대비 약 62% 수준으로 추정됨 (개당 환산 기준)"])
    asp.append(["※ 본 ASP 전망은 대외비이며 외부 공개 자료에 존재하지 않는 내부 추정치임"])
    style_header(asp, 6)
    for column, width in zip("ABCDEF", (12, 18, 16, 18, 24, 30)):
        asp.column_dimensions[column].width = width

    # -- 시트 3: 경쟁사 생산능력 (Core Gap 1 — 내부 추정) ---------------------
    capacity = workbook.create_sheet("경쟁사_생산능력_내부추정")
    capacity.append(["업체", "생산기지", "2025-Q4 월간 웨이퍼(장)", "2026-Q2 월간 웨이퍼(장)", "2026-Q4 월간 웨이퍼(장)", "HBM4 수율 추정(%)", "비고"])
    capacity_rows = [
        ["삼성전자(자사)", "평택 P4", 38_000, 50_000, 62_000, 68, "내부 실측 수율 — 외부 보도(50~60%)와 상이"],
        ["SK하이닉스", "이천 M16", 32_000, 40_000, 48_000, 82, "외부 보도 월 4만장과 일치"],
        ["SK하이닉스", "청주 M15X", 18_000, 26_000, 34_000, 80, "신규 라인 램프업"],
        ["Micron", "타이중 Fab", 14_000, 19_000, 25_000, 55, "내부 추정 점유율 15% 수준"],
    ]
    for row in capacity_rows:
        capacity.append(row)
    capacity.append([])
    capacity.append(["※ 내부 추정 기준 2026년 HBM4 점유율: SK하이닉스 62%, 자사 28%, Micron 15% 미만"])
    capacity.append(["※ 외부 일부 분석(점유율 70%/20%/22%)과 차이가 있음 — 자사 고객 인증 일정 반영 시 자사 점유율 상향 추정"])
    style_header(capacity, 7)
    for column, width in zip("ABCDEFG", (16, 14, 22, 22, 22, 18, 44)):
        capacity.column_dimensions[column].width = width

    workbook.save(path)

    top5 = sorted(by_customer.items(), key=lambda pair: pair[1], reverse=True)[:5]
    return {
        "file": path.name,
        "수주현황_row_count": len(order_rows),
        "총_수주액_백만달러": round(total_order_musd, 1),
        "고객사별_상위5": top5,
        "제품별_수주액": by_product,
        "HBM4_수주액_백만달러": round(hbm4_order_musd, 1),
    }


# ---------------------------------------------------------------------------
# ③ DOCX — DS부문 4Q 내부 결산 보고 (Core Gap 3 + 일치/불일치 수치)
# ---------------------------------------------------------------------------

def generate_settlement_docx(path: Path) -> dict:
    import docx
    from docx.shared import Pt, RGBColor

    document = docx.Document()

    def add_heading(text: str, level: int = 1) -> None:
        document.add_heading(text, level=level)

    def add_para(text: str, bold: bool = False) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold

    def add_confidential_banner() -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("■ 대외비 (CONFIDENTIAL) — 사외 반출 금지 ■")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    add_confidential_banner()
    document.add_heading("DS부문 2025년 4분기 내부 결산 보고", level=0)
    add_para("작성: DS부문 경영지원실 관리회계팀 | 보고일: 2026-01-22 | 문서번호: DS-FIN-2026-0118")

    add_heading("1. 전사 실적 요약", 1)
    add_para(
        "2025년 4분기 연결 기준 매출은 93.8조원, 영업이익은 20.1조원으로 분기 기준 역대 최대 실적을 기록하였다. "
        "DS부문은 전사 영업이익의 82%를 차지하며 실적 개선을 주도하였다."
    )

    add_heading("2. DS부문 실적 상세 (관리회계 기준)", 1)
    add_para(
        "DS부문 4분기 매출은 44조원으로 공시 기준과 동일하다. "
        "다만 내부 관리회계 기준 4분기 영업이익은 15.8조원으로 집계되었다. "
        "공시 기준(16.4조원)과의 차이 0.6조원은 본사 공통비 배부 기준 및 내부 이전가격 조정에 기인한다.",
        bold=True,
    )
    table = document.add_table(rows=5, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["구분", "매출(조원)", "영업이익(조원)", "비고"]
    data = [
        ["메모리", "33.4", "17.2", "HBM·DDR5 고부가 중심"],
        ["파운드리", "8.4", "-1.8", "선단공정 가동률 상승에도 적자 지속"],
        ["시스템LSI", "2.2", "-0.4", "엑시노스 수주 부진"],
        ["DS부문 합계", "44.0", "15.8 (관리회계) / 16.4 (공시)", "차이: 공통비 배부 기준"],
    ]
    for column_index, header_text in enumerate(headers):
        table.rows[0].cells[column_index].text = header_text
    for row_index, row_data in enumerate(data, start=1):
        for column_index, cell_text in enumerate(row_data):
            table.rows[row_index].cells[column_index].text = cell_text

    add_heading("3. 원가 구조 분석 (대외비 — 외부 미공개)", 1)
    add_para(
        "4분기 메모리 사업부 원가율은 58.3%로 전분기 대비 2.1%p 개선되었다. "
        "공정 단계별 비용 구성은 전공정 62%, 후공정(패키징) 23%, 테스트 15%이다. "
        "HBM4의 하이브리드 본딩 전환에 따라 2026년 후공정 비용 비중은 28%까지 상승할 것으로 전망된다."
    )
    add_para(
        "4분기 재고자산 평가충당금은 1.2조원이 신규 적립되었으며, "
        "이 중 0.8조원은 중국 수출 제재 영향에 따른 레거시 DRAM 재고, "
        "0.4조원은 파운드리 선단공정 시험생산 물량에 해당한다."
    )

    add_heading("4. 연구개발비 및 설비투자", 1)
    add_para(
        "내부 집계 기준 2025년 연간 연구개발비는 36.2조원이다. "
        "(주: 공시 기준 37.7조원과의 차이는 정부 과제 매칭 펀드 및 자본화 R&D 처리 기준 차이) "
        "2025년 DS부문 시설투자는 47.5조원으로 전사 설비투자의 약 90%를 차지하였다.",
        bold=True,
    )

    add_heading("5. 2026년 내부 경영 목표 (대외비)", 1)
    add_para("- DS부문 연간 매출 목표: 185조원 (전년 대비 +22%)")
    add_para("- HBM 매출 비중 목표: DS 메모리 매출의 35% (2025년 24% 대비 +11%p)")
    add_para("- HBM4 16단 양산 수율 목표: 2026년 2분기 75%, 4분기 82%")
    add_para("- 파운드리 손익분기 도달 목표: 2026년 4분기")
    add_para("- 평택 P4 라인 월간 웨이퍼 투입 목표: 2026년 2분기 50,000장 → 4분기 62,000장")

    add_heading("6. 리스크 요인", 1)
    add_para("- 환율: 원/달러 환율 10원 변동 시 분기 영업이익 약 0.3조원 변동")
    add_para("- 중국 제재: 레거시 DRAM 추가 충당금 최대 0.5조원 발생 가능")
    add_para("- HBM4 수율: 16단 적층 수율이 목표(75%) 대비 10%p 미달 시 2026년 매출 목표의 약 4% 하향 필요")

    document.save(path)
    return {"file": path.name, "주요_불일치_수치": ["DS 4Q 영업이익 15.8조 (웹 16.4조)", "연간 R&D 36.2조 (웹 37.7조)"]}


# ---------------------------------------------------------------------------
# ④ PDF — HBM4 경쟁구도 내부평가 (Core Gap 1 + 일치/불일치 수치)
#    python-docx로 생성 후 Word COM으로 PDF 변환
# ---------------------------------------------------------------------------

def generate_competition_pdf(pdf_path: Path) -> dict:
    import docx
    from docx.shared import Pt, RGBColor

    docx_path = pdf_path.with_suffix(".docx.tmp.docx")
    document = docx.Document()

    paragraph = document.add_paragraph()
    run = paragraph.add_run("■ 대외비 (CONFIDENTIAL) — 사외 반출 금지 ■")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    document.add_heading("HBM4 시장 경쟁구도 내부 평가 보고서", level=0)
    document.add_paragraph("작성: 메모리사업부 시장전략팀 | 보고일: 2026-02-10 | 문서번호: MEM-STR-2026-0034")

    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "2026년 HBM4 시장 규모는 546억 달러로 전년 대비 58% 성장이 전망된다. "
        "외부 분석기관 다수는 SK하이닉스 70%, 자사 20%, Micron 22% 수준의 점유율을 제시하고 있으나, "
        "당사 고객 인증 진행 상황과 수주 파이프라인을 반영한 내부 분석 기준으로는 "
        "SK하이닉스 62%, 자사 28%, Micron 15% 미만으로 추정된다. "
        "외부 추정과의 차이는 외부 기관이 당사의 엔비디아 루빈 2차 공급사 인증 일정을 반영하지 못한 데 기인한다."
    )

    document.add_heading("2. 시장 규모 및 기술 전환", level=1)
    document.add_paragraph(
        "HBM4는 2.8TB/s 대역폭, 11.7Gbps 핀 속도를 제공하며 HBM3E 대비 전력효율이 20% 개선된다. "
        "AI 가속기(엔비디아 루빈, AMD MI450) 수요 확대로 2026년 HBM 전체 시장에서 HBM4 비중은 "
        "하반기 기준 55%를 상회할 것으로 전망된다."
    )

    document.add_heading("3. 경쟁사 수율 비교 (내부 실측/추정)", level=1)
    table = document.add_table(rows=4, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["업체", "HBM4 12단 수율", "HBM4 16단 수율", "근거"]
    rows = [
        ["삼성전자(자사)", "68% (내부 실측)", "41% (내부 실측)", "평택 P4 양산라인 2026년 1월 실측 — 외부 보도(50~60%)는 과소평가"],
        ["SK하이닉스", "82% (추정)", "58% (추정)", "고객사 발주 패턴 역산 — 외부 보도(80~90%)와 대체로 일치"],
        ["Micron", "55% (추정)", "양산 전", "샘플 단계로 평가"],
    ]
    for column_index, header_text in enumerate(headers):
        table.rows[0].cells[column_index].text = header_text
    for row_index, row_data in enumerate(rows, start=1):
        for column_index, cell_text in enumerate(row_data):
            table.rows[row_index].cells[column_index].text = cell_text

    document.add_heading("4. 생산능력 비교 (Core Gap — 외부 자료 부재 영역)", level=1)
    document.add_paragraph(
        "당사 평택 P4 라인은 2026년 2분기 월 50,000장 웨이퍼 투입을 목표로 램프업 중이며, "
        "2026년 4분기 62,000장까지 확대 예정이다. "
        "SK하이닉스는 이천 M16에서 월 40,000장(외부 보도와 일치), 청주 M15X에서 월 26,000장 수준으로 파악된다. "
        "Micron 타이중 팹은 월 19,000장 수준으로 추정되어 3강 구도에서 캐파 열위가 지속될 전망이다."
    )

    document.add_heading("5. 고객 인증 진행 현황 (대외비)", level=1)
    document.add_paragraph("- 엔비디아 루빈(Rubin) 플랫폼: HBM4 12단 퀄리피케이션 테스트 2026년 1월 통과. 2차 공급사 지위 확보. 2026년 3분기 양산 공급 개시 예정.")
    document.add_paragraph("- AMD MI450: HBM4 12단/16단 동시 인증 진행 중. 2026년 2분기 완료 예상.")
    document.add_paragraph("- Google TPU v8: HBM4 16단 단독 공급 협상 중 (경쟁사 미참여). 성사 시 2027년 점유율 +5%p 효과.")
    document.add_paragraph("- Broadcom/Meta 커스텀 ASIC: 로직다이 커스터마이징 공동 개발 NDA 체결 (2026년 1월).")

    document.add_heading("6. 기술 격차 자체 평가", level=1)
    document.add_paragraph(
        "하이브리드 본딩 양산 적용 시점 기준 SK하이닉스 대비 약 6개월 격차가 존재하나, "
        "1c DRAM 공정 전환은 당사가 1분기 앞서 있어 2027년 HBM4E 세대에서 역전 가능성이 있다. "
        "로직다이 커스터마이징 역량은 자사 파운드리 내재화로 TSMC 의존적인 경쟁사 대비 구조적 우위로 평가된다."
    )

    document.add_heading("7. 전략 권고", level=1)
    document.add_paragraph("- 엔비디아 루빈 2차 공급 물량의 조기 확대를 위해 16단 수율 개선 TF에 분기 0.3조원 추가 투입 권고")
    document.add_paragraph("- Google TPU v8 단독 공급 협상의 우선순위 상향 — 경쟁사 미참여 상태인 현 시점이 협상 적기")
    document.add_paragraph("- 외부 시장조사기관 대상 자사 점유율 전망(28%) 커뮤니케이션 강화 필요")

    document.save(docx_path)

    # Word COM으로 PDF 변환
    converted = _convert_docx_to_pdf(docx_path, pdf_path)
    if converted:
        docx_path.unlink(missing_ok=True)
        return {"file": pdf_path.name, "변환": "Word COM PDF 변환 성공"}
    # 변환 실패 시 docx로 대체 (이름 변경)
    fallback = pdf_path.with_suffix(".docx")
    docx_path.rename(fallback)
    return {"file": fallback.name, "변환": "PDF 변환 실패 — docx로 대체됨"}


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        import win32com.client
        import pythoncom

        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(docx_path.resolve()))
            doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)  # 17 = wdFormatPDF
            doc.Close(False)
            return True
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
    except Exception as exc:  # noqa: BLE001
        print(f"[warn] Word COM PDF 변환 실패: {exc}")
        return False


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main() -> None:
    _check_deps()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    EXPECTED_DIR.mkdir(parents=True, exist_ok=True)

    print(f"출력 위치: {OUTPUT_DIR}")
    expected: dict = {"생성일": "2026-06-02", "시나리오": "A — 삼성전자 DS/HBM"}

    print("[1/4] CSV 생성 중 (DS부문 월별 사업부 실적, 480행)...")
    expected["csv"] = generate_csv(OUTPUT_DIR / "DS부문_월별_사업부실적_2025.csv")

    print("[2/4] XLSX 생성 중 (HBM 고객사별 수주현황, 3개 시트)...")
    expected["xlsx"] = generate_xlsx(OUTPUT_DIR / "HBM_고객사별_수주현황_2026_대외비.xlsx")

    print("[3/4] DOCX 생성 중 (DS부문 4Q 내부 결산 보고)...")
    expected["docx"] = generate_settlement_docx(OUTPUT_DIR / "DS부문_2025년_4분기_내부결산보고_대외비.docx")

    print("[4/4] PDF 생성 중 (HBM4 경쟁구도 내부 평가 / Word COM 변환)...")
    expected["pdf"] = generate_competition_pdf(OUTPUT_DIR / "HBM4_경쟁구도_내부평가_대외비.pdf")

    # crosscheck 기대값 요약
    expected["crosscheck_설계"] = {
        "일치(supports) 예상": [
            "4Q25 연결 매출 93.8조원 / 영업이익 20.1조원",
            "DS부문 4Q 매출 44조원",
            "DS부문 이익 비중 82%",
            "HBM4 시장 규모 546억 달러 / 성장률 58%",
            "HBM4 대역폭 2.8TB/s",
            "DS 설비투자 47.5조원",
            "SK하이닉스 M16 월 4만장",
        ],
        "불일치(numeric_mismatch) 예상": [
            "DS 4Q 영업이익: 내부 15.8조원 vs 웹 16.4조원",
            "연간 R&D: 내부 36.2조원 vs 웹 37.7조원",
            "SK하이닉스 HBM4 점유율: 내부 62% vs 웹 70%",
            "자사 HBM4 수율: 내부 68% vs 웹 보도 50~60%",
            "Micron 점유율: 내부 15% vs 웹 22%",
        ],
        "내부 전용(Core Gap 채움)": [
            "Gap1: 3사 생산능력(웨이퍼/월)·고객 인증 현황·기술 격차 평가",
            "Gap2: HBM4 vs HBM3E 분기별 ASP·가격 프리미엄(62%)",
            "Gap3: 원가율 58.3%·공정별 비용·재고충당금 1.2조원 상세",
        ],
    }

    expected_path = EXPECTED_DIR / "expected_values.json"
    expected_path.write_text(
        json.dumps(expected, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print("\n=== 생성 완료 ===")
    for item in OUTPUT_DIR.iterdir():
        print(f"  {item.name}  ({item.stat().st_size:,} bytes)")
    print(f"\n검산 기대값(테스터용, 로컬 폴더에 등록 금지): {expected_path}")


if __name__ == "__main__":
    main()
